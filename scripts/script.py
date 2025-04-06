import os
import json
from crewai import Agent, Task, Crew, Process
from langchain_upstage import ChatUpstage
from crewai_tools import PDFSearchTool
from openai import OpenAI
import gradio as gr
import re
import json
from fpdf import FPDF
from docx import Document
import gradio as gr
from docx.shared import Pt

# Set environment variables
os.environ["UPSTAGE_API_BASE"] = "https://api.upstage.ai/v1/solar"
os.environ["UPSTAGE_API_KEY"] = "up_sxQRRcrbTmgfXVNaxuWpTgkd5Yuig"

class ScriptGenerator:
    def __init__(self, characters_num,cafe_name,cafe_environment):
        self.characters_num = characters_num
        self.cafe_name= cafe_name
        self.cafe_environment = cafe_environment
        self.llm = ChatUpstage()
        self.rag_tool = PDFSearchTool(
            pdf='/Users/debbiechoonghuitian/Jejom-1/scripts/outputs/Jeju.pdf',
            config=dict(
                llm=dict(
                    provider="openai",
                    config=dict(
                        model="solar-1-mini-chat",
                    ),
                ),
                embedder=dict(
                    provider="huggingface",
                    config=dict(
                        model="BAAI/bge-small-en-v1.5",
                    ),
                ),
            )
        )
        self.setup_agents_and_tasks()

    def setup_agents_and_tasks(self):
        # Define the Script Planner Agent
        self.script_planner = Agent(
            role="Script Planner",
            goal=f"Design a complex and engaging {self.characters_num}-character storyline inspired by the myths and legends of Jeju Island, Korea, to bring players challenge and enjoyment.",
            backstory="""You are an experienced storyteller, well-versed in the essence of humanity and drama. You excel at constructing tight plots that immerse players.""",
            verbose=True,
            allow_delegation=False,
            llm=self.llm
        )
        self.titler = Agent(
            role="Title",
            goal=f"Design a title for the murder mystery game",
            backstory="""You are an experienced storyteller, well-versed in the essence of humanity and drama. You excel at constructing tight plots that immerse players.""",
            verbose=True,
            allow_delegation=False,
            llm=self.llm
        )
        self.timer = Agent(
            role="Duration",
            goal=f"Write the duration of the time.",
            backstory="""You are an experienced storyteller, well-versed in the essence of humanity and drama. You excel at constructing tight plots that immerse players.""",
            verbose=True,
            allow_delegation=False,
            llm=self.llm
        )

        # Define the Character Designer Agent
        self.character_designer = Agent(
            role="Character Designer",
            goal="Create unique and detailed backstories, motivations, and secrets for each character, bringing them to life.",
            backstory="""You are an expert in character design, adept at creating complex and believable character profiles. Your designs will add depth to the story.""",
            verbose=True,
            allow_delegation=False,
            llm=self.llm
        )

        # Define the Script Writer Agent
        self.script_writer_agent = Agent(
            role="Script Writer",
            goal="Write the script of each character to explain to the player about the character's life and the events that lead to the main crime scene.",
            backstory="""A playwright by trade, you have a deep understanding of the character's role in the murder mystery game.""",
            verbose=True,
            allow_delegation=False,
            llm=self.llm
        )

        # Define the Player Instruction Writer Agent
        self.player_writer_agent = Agent(
            role="Player Instruction Writer",
            goal="Create detailed and engaging roleplay instructions for each character, focusing on the five rounds of discussions described in the task.",
            backstory="""Once a renowned playwright who brought life to the most intricate and suspenseful mysteries on stage, this agent has transitioned to the world of digital narratives. Their experience in crafting compelling character arcs and dramatic plot twists is unmatched. Having grown up in a family of storytellers in a quaint, coastal village known for its legends and lore, the agent carries the legacy of weaving tales that captivate audiences. Now, they apply their expertise to create immersive roleplaying experiences, ensuring every character is memorable, every line meaningful, and every plot twist a revelation.""",
            verbose=True,
            allow_delegation=False,
            llm=self.llm
        )

        # Define the Clue Generator Agent
        self.clue_generator = Agent(
            role="Clue Generator",
            goal="Design challenging and misleading clues to guide players through the game and solve the mystery.",
            backstory="""You are a puzzle enthusiast, skilled at designing complex and deceptive clues that make the game more challenging.""",
            verbose=True,
            allow_delegation=False,
            llm=self.llm
        )

        self.define_tasks()
    
    def define_tasks(self):
        self.background_setting_task = Task(
            description="""Determine the background setting for the murder mystery, drawing inspiration from the legends, myths, history, and culture of Jeju Island. The storyline does not necessarily need to be set on Jeju Island. Focus on creating a rich, atmospheric setting that influences the characters' actions and motivations. The story should revolve around four key characters, each connected to these cultural and historical elements in some way.""",
            expected_output="A detailed background story inspired by Jeju Island's legends, history, and culture, including the crime scene setting, the course of the case, and key event descriptions. ",
            agent=self.script_planner,
            output_file="background_setting.txt"
        )

        self.character_creation_task = Task(
            description="""Create complete character profiles for each character, including backstories, motivations, secrets, and relationships with other characters. Ensure that each character has a unique personality and story.""",
            expected_output="A detailed list of all characters, with each character having a complete profile, including backstories, motivations, secrets, and relationships.",
            agent=self.character_designer,
            output_file="character.txt"
        )

        self.script_writing_task = Task(
            description=f"""Write a complete 4-day event log leading to the crime day for all the {self.characters_num} characters. Ensure each log is fully written, with no unfinished sentences or thoughts. Include specific dates, key events, thoughts, plans, and interactions that provide insight into the character's motives and actions. Each event log should conclude with a summary or reflective thought that naturally completes the narrative.""",
            expected_output="A complete 4-day event log for each character. Write for all characters",
            agent=self.script_writer_agent,
            output_file="character_event_log.txt"
        )

        self.player_writing_instruction_task = Task(
            description="""There will be 5 rounds of role play discussions. The first round will have players introduce themselves and discuss the crime. The second round is to collect all the clues from the other players. The third round is an open discussion. The fourth round is to collect clues. The fifth round is to vote out on the murderer.

            Create a guide and explanation on how to roleplay each character in these rounds. For each character, write the character's name followed by the detailed instructions for each round of discussion for that character. Don't expose the murderer. Put all the instructions for each character together in a clear and organized format.""",
            expected_output="A list of guide and explanation on what the players have to do to roleplay their characters, titled with the character's name followed by detailed instructions for each round. Ensure that each character's instructions are presented together and clearly labeled. Make it more personal for each character.",
            agent=self.player_writer_agent,
            output_file="player_instructions.txt"
        )

        self.clue_design_task = Task(
            description="""Design 4 clues that players discover in the game, ensuring that these clues are challenging and fit the plot logic. Each character should have 2 key clues and 2 misleading clues. Title each section with the character's name followed by the list of their clues and misleading clues.""",
            expected_output="A list of 4 clues for each character, including 2 key clues and 2 misleading clues, with explanations of their role in the story. Ensure each character's section is titled with the character's name, and the clues are clearly labeled as key clues or misleading clues.",
            agent=self.clue_generator,
            output_file="player_clues.txt"
        )
        self.title = Task(
            description="""Write the title for the script generated.""",
            expected_output="Output a title",
            agent=self.titler,
            output_file="title.txt"
        )
        self.time = Task(
            description="""Write the game duraction""",
            expected_output=" Just right the answer as '2-3hours'. Don't write full sentences, just a direct duration",
            agent=self.timer,
            output_file="time_taken.txt"
        )

    def run_tasks(self):
        crew = Crew(
            agents=[self.script_planner, self.character_designer, self.script_writer_agent, self.clue_generator, self.player_writer_agent, self.titler,self.timer],
            tasks=[self.background_setting_task, self.character_creation_task, self.script_writing_task, self.clue_design_task, self.player_writing_instruction_task, self.title,self.time],
            verbose=True,
            process=Process.sequential
        )

        crew.kickoff()
        task_outputs = {}

        for task in crew.tasks:
            try:
                with open(task.output_file, 'r') as f:
                    task_outputs[task.agent.role] = f.read()
                print(f"Successfully read output for {task.agent.role}")
            except FileNotFoundError:
                print(f"File not found for {task.agent.role}: {task.output_file}")
            except Exception as e:
                print(f"Error reading file for {task.agent.role}: {e}")

        output_json_path = "all_task_outputs.json"
        with open(output_json_path, 'w') as json_file:
            json.dump(task_outputs, json_file, indent=5)
        
        print("All task outputs have been saved to:", output_json_path)

        return output_json_path, self.cafe_name



import json

class Translator:
    def __init__(self, api_key):
        self.client = OpenAI(api_key=api_key, base_url="https://api.upstage.ai/v1/solar")

    def translate_text(self, text, model="solar-1-mini-translate-enko", chunk_size=1600):
        # Split the text into chunks, making sure not to split in the middle of a sentence
        import re
        
        # Split text into chunks, trying to avoid splitting sentences
        def split_text(text, chunk_size):
            sentences = re.split(r'(?<=[.!?]) +', text)
            chunks, current_chunk = [], ""
            for sentence in sentences:
                if len(current_chunk) + len(sentence) + 1 <= chunk_size:
                    if current_chunk:
                        current_chunk += " " + sentence
                    else:
                        current_chunk = sentence
                else:
                    chunks.append(current_chunk)
                    current_chunk = sentence
            if current_chunk:
                chunks.append(current_chunk)
            return chunks

        chunks = split_text(text, chunk_size)
        translated_text = ""

        for chunk in chunks:
            stream = self.client.chat.completions.create(
                model=model,
                messages=[
                    {
                        "role": "assistant",
                        "content": chunk
                    }
                ],
                stream=True,
            )
            for chunk in stream:
                if chunk.choices[0].delta.content:
                    translated_text += chunk.choices[0].delta.content

        return translated_text


    def translate_and_save(self, input_file, output_file):
        with open(input_file, 'r', encoding='utf-8') as file:
            data = json.load(file)

        translated_data = {}
        for key, text in data.items():
            print(f"Translating {key}...")
            translated_text = self.translate_text(text)
            print(translated_text)
            translated_data[key] = translated_text

        with open(output_file, 'w', encoding='utf-8') as file:
            json.dump(translated_data, file, ensure_ascii=False, indent=4)

        return output_file

        print(f"Translation complete. Check '{output_file}' for results.")



def create_docx(json_data,docx_name,cafe_name):
    docx_path = f"{docx_name}.docx"
    doc = Document()

    
    heading = doc.add_heading("Murder Mystery Game Script", level=1)
    run = heading.runs[0]
    run.font.size = Pt(24)
    doc.add_heading(cafe_name, level=1)

    # if 'title' in json_data:
    #     doc.add_paragraph("Title")
    #     doc.add_paragraph(json_data['title'])
    # if 'Duration' in json_data:
    #     doc.add_paragraph("Duration")
    #     doc.add_paragraph(json_data['Duration'])
    # if 'Script Planner' in json_data:
    #     doc.add_paragraph("Storyline")
    #     doc.add_paragraph(json_data['Script Planner'])
    # if 'Character Designer' in json_data:
    #     doc.add_paragraph("Characters")
    #     doc.add_paragraph(json_data['Character Designer'])
    # if 'Script Writer' in json_data:
    #     doc.add_paragraph("Events before the crime scene")
    #     doc.add_paragraph(json_data['Script Writer'])
    # if 'Clue Generator' in json_data:
    #     doc.add_paragraph("Clues")
    #     doc.add_paragraph(json_data['Clue Generator'])
    # if 'Player Instruction Writer' in json_data:
    #     doc.add_paragraph("Player Instructions")
    #     doc.add_paragraph(json_data['Player Instruction Writer'])

    if docx_name=='english_script':
        if 'Title' in json_data:
            title_paragraph = doc.add_paragraph("Title")
            run = title_paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(14)  # Adjust the font size as needed
            doc.add_paragraph(json_data['Title'])

        if 'Duration' in json_data:
            duration_paragraph = doc.add_paragraph("Duration")
            run = duration_paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            doc.add_paragraph(json_data['Duration'])

        if 'Script Planner' in json_data:
            storyline_paragraph = doc.add_paragraph("Storyline")
            run = storyline_paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            doc.add_paragraph(json_data['Script Planner'])

        if 'Character Designer' in json_data:
            characters_paragraph = doc.add_paragraph("Characters")
            run = characters_paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            doc.add_paragraph(json_data['Character Designer'])

        if 'Script Writer' in json_data:
            events_paragraph = doc.add_paragraph("Events before the crime scene")
            run = events_paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            doc.add_paragraph(json_data['Script Writer'])

        if 'Clue Generator' in json_data:
            clues_paragraph = doc.add_paragraph("Clues")
            run = clues_paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            doc.add_paragraph(json_data['Clue Generator'])

        if 'Player Instruction Writer' in json_data:
            instructions_paragraph = doc.add_paragraph("Player Instructions")
            run = instructions_paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            doc.add_paragraph(json_data['Player Instruction Writer'])

    elif docx_name=='korean_script':
        if 'Title' in json_data:
            title_paragraph = doc.add_paragraph("제목")
            run = title_paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(14)  # Adjust the font size as needed
            doc.add_paragraph(json_data['Title'])

        if 'Duration' in json_data:
            duration_paragraph = doc.add_paragraph("지속")
            run = duration_paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            doc.add_paragraph(json_data['Duration'])

        if 'Script Planner' in json_data:
            storyline_paragraph = doc.add_paragraph("줄거리")
            run = storyline_paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            doc.add_paragraph(json_data['Script Planner'])

        if 'Character Designer' in json_data:
            characters_paragraph = doc.add_paragraph("성격")
            run = characters_paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            doc.add_paragraph(json_data['Character Designer'])

        if 'Script Writer' in json_data:
            events_paragraph = doc.add_paragraph("범죄 현장 이전의 사건")
            run = events_paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            doc.add_paragraph(json_data['Script Writer'])

        if 'Clue Generator' in json_data:
            clues_paragraph = doc.add_paragraph("단서")
            run = clues_paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            doc.add_paragraph(json_data['Clue Generator'])

        if 'Player Instruction Writer' in json_data:
            instructions_paragraph = doc.add_paragraph("플레이어 지침")
            run = instructions_paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            doc.add_paragraph(json_data['Player Instruction Writer'])

    # for key, value in json_data.items():
    #     doc.add_paragraph(f"{key}: {value}")
    
    doc.save(docx_path)
    return docx_path

#generate riles 
def generate_files(original_script,translated_script,cafe_name):
    english_script = create_docx(original_script,"english_script",cafe_name)
    korean_script = create_docx(translated_script,"korean_script",cafe_name)
    return english_script, korean_script



def generate_scripts(characters_num,cafe_name,cafe_environment):
    script_generator = ScriptGenerator(characters_num,cafe_name,cafe_environment)
    output_json_path, cafe_name = script_generator.run_tasks()
    translator = Translator(api_key="up_sxQRRcrbTmgfXVNaxuWpTgkd5Yuig")
    translator_json_path = translator.translate_and_save(input_file=output_json_path, output_file='translated_file.json')

    with open(output_json_path, 'r', encoding='utf-8') as file:
        original_script = json.load(file)

    with open(translator_json_path, 'r', encoding='utf-8') as file:
        translated_script = json.load(file)
        
    return generate_files(original_script,translated_script,cafe_name)





    
    
with gr.Blocks() as demo:
    gr.Markdown(
    """
    # Create Your Own Murder Mystery Script (剧本杀)

    Begin crafting your story for your cafe below.

    """)


    cafe_name = gr.Textbox(label="What is the name of your cafe?")
    characters_num = gr.Textbox(label="How many characters would you like to feature in your murder mystery script?")
    cafe_environment= gr.Textbox(label="Briefly describe the ambience of your cafe",info="My cafe is a cozy, warm haven filled with natural light")

    btn = gr.Button("Generate Script")

    english_output = gr.File(label="Download English Version")
    korean_output = gr.File(label="Download Korean Version")
    
    btn.click(fn=generate_scripts, inputs=[characters_num,cafe_name,cafe_environment],outputs=[english_output, korean_output])

if __name__ == "__main__":
    demo.launch()
        
